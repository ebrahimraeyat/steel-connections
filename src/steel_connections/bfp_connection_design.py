"""
bfp_connection_design.py
ماژول طراحی اتصال پیچی با جفت سپری (BFP - Bolted Flange Plate connection)
بر اساس مبحث دهم مقررات ملی ساختمان (ویرایش ۱۴۰۱) - بخش اتصالات لرزه‌ای

این ماژول شامل:
- طراحی اتصال BFP برای قاب‌های خمشی ویژه (SMF)
- محاسبه نیروهای وارده بر پیچ‌ها
- طراحی ورق‌های سپری (بالایی و پایینی)
- طراحی جوش‌های اتصال ورق سپری به تیر
- کنترل ضوابط لرزه‌ای
- محاسبه ضخامت ورق و تعداد پیچ‌های موردنیاز
- خروجی به فرمت Word (docx)
"""

from enum import Enum
from dataclasses import dataclass
from typing import Optional, Tuple, List, Dict, Any
import math
from datetime import datetime
import os

# تلاش برای import کتابخانه python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ هشدار: کتابخانه python-docx نصب نیست. برای خروجی Word ابتدا نصب کنید: pip install python-docx")


# ============================================================================
# Enum‌های تعریف انواع
# ============================================================================

class BoltType(Enum):
    """نوع پیچ پرمقاومت"""
    A325 = "A325"
    A490 = "A490"


class DesignMethod(Enum):
    """روش طراحی"""
    LRFD = "LRFD"
    ASD = "ASD"


class ConnectionType(Enum):
    """نوع اتصال BFP"""
    TOP_FLANGE_ONLY = "top_flange_only"          # فقط ورق سپری بالایی
    BOTTOM_FLANGE_ONLY = "bottom_flange_only"    # فقط ورق سپری پایینی
    BOTH_FLANGES = "both_flanges"                # هر دو ورق سپری


# ============================================================================
# کلاس‌های داده
# ============================================================================

@dataclass
class BoltGroupResult:
    """نتایج طراحی گروه پیچ"""
    bolt_diameter_mm: float
    bolt_type: BoltType
    num_bolts: int
    bolt_force_n: float
    bolt_capacity_n: float
    is_adequate: bool
    utilization_ratio: float


@dataclass
class FlangePlateResult:
    """نتایج طراحی ورق سپری"""
    plate_thickness_mm: float
    plate_width_mm: float
    plate_length_mm: float
    required_thickness_mm: float
    yield_check_ratio: float
    rupture_check_ratio: float
    block_shear_ratio: float
    is_adequate: bool


@dataclass
class WeldResult:
    """نتایج طراحی جوش"""
    weld_size_mm: float
    weld_length_mm: float
    required_size_mm: float
    capacity_n: float
    utilization_ratio: float
    is_adequate: bool


@dataclass
class BFPConnectionDesign:
    """نتایج نهایی طراحی اتصال BFP"""
    # اطلاعات کلی
    connection_type: ConnectionType
    design_method: DesignMethod
    
    # مشخصات تیر
    fy_beam_mpa: float
    fu_beam_mpa: float
    zx_beam_mm3: float
    beam_depth_mm: float
    beam_flange_width_mm: float
    beam_flange_thickness_mm: float
    
    # مشخصات ورق
    plate_fy_mpa: float
    plate_fu_mpa: float
    
    # نتایج محاسبات اصلی
    m_pr_nmm: float
    flange_force_n: float
    
    # نتایج طراحی پیچ‌ها
    top_bolts: Optional[BoltGroupResult]
    bottom_bolts: Optional[BoltGroupResult]
    
    # نتایج طراحی ورق‌های سپری
    top_plate: Optional[FlangePlateResult]
    bottom_plate: Optional[FlangePlateResult]
    
    # نتایج طراحی جوش‌ها
    top_weld: Optional[WeldResult]
    bottom_weld: Optional[WeldResult]
    
    # کنترل‌های لرزه‌ای
    seismic_check_passed: bool
    seismic_checks: List[str]
    overstrength_ratio: float
    
    # جمع‌بندی
    is_valid: bool
    summary: str
    design_date: str = None
    
    def __post_init__(self):
        if self.design_date is None:
            self.design_date = datetime.now().strftime("%Y/%m/%d - %H:%M:%S")


# ============================================================================
# توابع کمکی داخلی
# ============================================================================

def _get_bolt_area(diameter_mm: float) -> float:
    """محاسبه مساحت اسمی پیچ (mm²)"""
    return math.pi * (diameter_mm ** 2) / 4.0


def _get_bolt_tension_capacity(bolt_type: BoltType, diameter_mm: float, design_method: DesignMethod) -> float:
    """محاسبه ظرفیت کششی اسمی یک پیچ"""
    A_b = _get_bolt_area(diameter_mm)
    
    if bolt_type == BoltType.A325:
        F_nt = 620.0  # MPa
    else:
        F_nt = 780.0  # MPa
    
    R_n = F_nt * A_b  # N
    
    if design_method == DesignMethod.LRFD:
        phi = 0.75
        capacity = phi * R_n
    else:
        omega = 2.00
        capacity = R_n / omega
    
    return capacity


def _get_plate_yield_capacity(thickness_mm: float, width_mm: float, fy_mpa: float, design_method: DesignMethod) -> float:
    """محاسبه ظرفیت تسلیم کششی ورق"""
    A_g = thickness_mm * width_mm
    R_n = fy_mpa * A_g
    
    if design_method == DesignMethod.LRFD:
        phi = 0.90
        capacity = phi * R_n
    else:
        omega = 1.67
        capacity = R_n / omega
    
    return capacity


def _get_plate_rupture_capacity(thickness_mm: float, width_mm: float, fu_mpa: float,
                                 bolt_diameter_mm: float, num_bolts_in_row: int, design_method: DesignMethod) -> float:
    """محاسبه ظرفیت گسیختگی کششی ورق با در نظر گرفتن سوراخ‌ها"""
    A_g = thickness_mm * width_mm
    hole_diameter = bolt_diameter_mm + 2.0
    
    A_n = A_g - num_bolts_in_row * hole_diameter * thickness_mm
    U = 0.85
    A_e = U * A_n
    
    R_n = fu_mpa * A_e
    
    if design_method == DesignMethod.LRFD:
        phi = 0.75
        capacity = phi * R_n
    else:
        omega = 2.00
        capacity = R_n / omega
    
    return capacity


def _check_block_shear(thickness_mm: float, width_mm: float, fy_mpa: float, fu_mpa: float,
                       bolt_diameter_mm: float, edge_distance_mm: float, bolt_spacing_mm: float,
                       num_bolts_in_row: int, design_method: DesignMethod) -> Tuple[float, float]:
    """
    بررسی بلوک برشی در ورق سپری
    
    Returns:
        (capacity, utilization_ratio)
    """
    hole_diameter = bolt_diameter_mm + 2.0
    
    shear_length = edge_distance_mm + (num_bolts_in_row - 1) * bolt_spacing_mm
    A_gv = 2 * shear_length * thickness_mm
    A_nv = 2 * (shear_length - num_bolts_in_row * hole_diameter) * thickness_mm
    
    tension_length = bolt_spacing_mm
    A_nt = tension_length * thickness_mm - hole_diameter * thickness_mm
    
    R_n1 = 0.6 * fu_mpa * A_nv + fu_mpa * A_nt
    R_n2 = 0.6 * fy_mpa * A_gv + fu_mpa * A_nt
    
    R_n = min(R_n1, R_n2)
    
    if design_method == DesignMethod.LRFD:
        phi = 0.75
        capacity = phi * R_n
    else:
        omega = 2.00
        capacity = R_n / omega
    
    return capacity, capacity  # برای سازگاری، capacity دوم برای نسبت مصرف بعداً محاسبه می‌شود


# ============================================================================
# تابع اصلی طراحی اتصال BFP
# ============================================================================

def design_bfp_connection(
    # مشخصات تیر
    fy_beam_mpa: float,
    fu_beam_mpa: float,
    zx_beam_mm3: float,
    beam_depth_mm: float,
    beam_flange_width_mm: float,
    beam_flange_thickness_mm: float,
    
    # مشخصات پیچ
    bolt_type: BoltType = BoltType.A325,
    bolt_diameter_mm: float = 20.0,
    bolt_grade_ry: float = 1.1,
    
    # مشخصات ورق سپری
    plate_fy_mpa: float = 345.0,
    plate_fu_mpa: float = 450.0,
    
    # پارامترهای هندسی
    edge_distance_mm: float = 40.0,
    bolt_spacing_mm: float = 80.0,
    
    # پارامترهای جوش
    electrode_fexx_mpa: float = 480.0,
    weld_size_mm: Optional[float] = None,
    
    # روش طراحی
    design_method: DesignMethod = DesignMethod.LRFD,
    connection_type: ConnectionType = ConnectionType.BOTH_FLANGES,
    
    # محدودیت‌ها
    max_num_bolts: int = 12,
    min_plate_thickness_mm: float = 10.0,
    max_plate_thickness_mm: float = 50.0,
    
) -> BFPConnectionDesign:
    """
    طراحی کامل اتصال پیچی با جفت سپری (BFP)
    """
    
    seismic_checks = []
    seismic_check_passed = True
    
    # ========================================================================
    # مرحله 1: محاسبه لنگر پلاستیک مورد انتظار (M_pr)
    # ========================================================================
    
    C_pr = min((fy_beam_mpa + fu_beam_mpa) / (2.0 * fy_beam_mpa), 1.2)
    M_pr = C_pr * bolt_grade_ry * fy_beam_mpa * zx_beam_mm3
    
    seismic_checks.append(f"M_pr = {M_pr/1e6:.1f} kN·m (C_pr={C_pr:.3f}, R_y={bolt_grade_ry})")
    
    # ========================================================================
    # مرحله 2: محاسبه نیروی کششی وارد بر ورق سپری
    # ========================================================================
    
    flange_force_n = M_pr / beam_depth_mm
    
    seismic_checks.append(f"نیروی کششی در هر ورق سپری = {flange_force_n/1000:.1f} kN")
    
    # ========================================================================
    # مرحله 3: طراحی پیچ‌ها
    # ========================================================================
    
    bolt_capacity_tension = _get_bolt_tension_capacity(bolt_type, bolt_diameter_mm, design_method)
    num_bolts_required = max(2, int(math.ceil(flange_force_n / bolt_capacity_tension)))
    
    if num_bolts_required > max_num_bolts:
        seismic_checks.append(f"⚠️ تعداد پیچ موردنیاز ({num_bolts_required}) بیشتر از حداکثر مجاز ({max_num_bolts}) است")
        seismic_check_passed = False
        num_bolts_required = max_num_bolts
    
    bolt_force_per_bolt = flange_force_n / num_bolts_required
    
    is_bolt_adequate = bolt_force_per_bolt <= bolt_capacity_tension
    bolt_utilization = bolt_force_per_bolt / bolt_capacity_tension if bolt_capacity_tension > 0 else 1.0
    
    top_bolts_result = BoltGroupResult(
        bolt_diameter_mm=bolt_diameter_mm,
        bolt_type=bolt_type,
        num_bolts=num_bolts_required,
        bolt_force_n=bolt_force_per_bolt,
        bolt_capacity_n=bolt_capacity_tension,
        is_adequate=is_bolt_adequate,
        utilization_ratio=bolt_utilization
    )
    
    bottom_bolts_result = BoltGroupResult(
        bolt_diameter_mm=bolt_diameter_mm,
        bolt_type=bolt_type,
        num_bolts=num_bolts_required,
        bolt_force_n=bolt_force_per_bolt,
        bolt_capacity_n=bolt_capacity_tension,
        is_adequate=is_bolt_adequate,
        utilization_ratio=bolt_utilization
    ) if connection_type in [ConnectionType.BOTH_FLANGES, ConnectionType.BOTTOM_FLANGE_ONLY] else None
    
    seismic_checks.append(f"پیچ‌ها: {num_bolts_required} عدد پیچ {bolt_type.value} قطر {bolt_diameter_mm}mm")
    seismic_checks.append(f"  نیروی هر پیچ = {bolt_force_per_bolt/1000:.1f} kN (ظرفیت: {bolt_capacity_tension/1000:.1f} kN)")
    seismic_checks.append(f"  نسبت مصرف = {bolt_utilization:.2f} {'✓' if is_bolt_adequate else '✗'}")
    
    # ========================================================================
    # مرحله 4: طراحی ورق سپری
    # ========================================================================
    
    plate_width = max(beam_flange_width_mm, min(beam_flange_width_mm * 1.5, beam_flange_width_mm + 80.0))
    
    # تسلیم
    if design_method == DesignMethod.LRFD:
        phi_yield = 0.90
        required_area_yield = flange_force_n / (phi_yield * plate_fy_mpa)
    else:
        omega_yield = 1.67
        required_area_yield = flange_force_n * omega_yield / plate_fy_mpa
    
    thickness_yield = required_area_yield / plate_width
    
    # گسیختگی
    hole_diameter = bolt_diameter_mm + 2.0
    A_n_net = plate_width - num_bolts_required * hole_diameter
    A_e = 0.85 * A_n_net
    
    if design_method == DesignMethod.LRFD:
        phi_rupture = 0.75
        required_area_rupture = flange_force_n / (phi_rupture * plate_fu_mpa)
    else:
        omega_rupture = 2.00
        required_area_rupture = flange_force_n * omega_rupture / plate_fu_mpa
    
    thickness_rupture = required_area_rupture / plate_width
    
    required_thickness = max(thickness_yield, thickness_rupture, min_plate_thickness_mm)
    required_thickness = max(required_thickness, beam_flange_thickness_mm, 10.0)
    
    standard_thicknesses = [10, 12, 15, 18, 20, 22, 25, 28, 30, 32, 35, 38, 40, 45, 50]
    plate_thickness = min([t for t in standard_thicknesses if t >= required_thickness], default=required_thickness)
    
    if plate_thickness > max_plate_thickness_mm:
        seismic_checks.append(f"⚠️ ضخامت ورق موردنیاز ({plate_thickness:.1f}mm) بیشتر از حداکثر مجاز ({max_plate_thickness_mm}mm) است")
        seismic_check_passed = False
        plate_thickness = max_plate_thickness_mm
    
    actual_capacity_yield = _get_plate_yield_capacity(plate_thickness, plate_width, plate_fy_mpa, design_method)
    yield_ratio = flange_force_n / actual_capacity_yield if actual_capacity_yield > 0 else 1.0
    
    actual_capacity_rupture = _get_plate_rupture_capacity(plate_thickness, plate_width, plate_fu_mpa,
                                                           bolt_diameter_mm, num_bolts_required, design_method)
    rupture_ratio = flange_force_n / actual_capacity_rupture if actual_capacity_rupture > 0 else 1.0
    
    block_shear_capacity, _ = _check_block_shear(
        plate_thickness, plate_width, plate_fy_mpa, plate_fu_mpa,
        bolt_diameter_mm, edge_distance_mm, bolt_spacing_mm, num_bolts_required, design_method
    )
    block_shear_ratio = flange_force_n / block_shear_capacity if block_shear_capacity > 0 else 1.0
    
    is_plate_adequate = (yield_ratio <= 1.0) and (rupture_ratio <= 1.0) and (block_shear_ratio <= 1.0)
    
    plate_length = 2 * edge_distance_mm + (num_bolts_required - 1) * bolt_spacing_mm + 50.0
    
    top_plate_result = FlangePlateResult(
        plate_thickness_mm=plate_thickness,
        plate_width_mm=plate_width,
        plate_length_mm=plate_length,
        required_thickness_mm=required_thickness,
        yield_check_ratio=yield_ratio,
        rupture_check_ratio=rupture_ratio,
        block_shear_ratio=block_shear_ratio,
        is_adequate=is_plate_adequate
    )
    
    bottom_plate_result = FlangePlateResult(
        plate_thickness_mm=plate_thickness,
        plate_width_mm=plate_width,
        plate_length_mm=plate_length,
        required_thickness_mm=required_thickness,
        yield_check_ratio=yield_ratio,
        rupture_check_ratio=rupture_ratio,
        block_shear_ratio=block_shear_ratio,
        is_adequate=is_plate_adequate
    ) if connection_type in [ConnectionType.BOTH_FLANGES, ConnectionType.BOTTOM_FLANGE_ONLY] else None
    
    seismic_checks.append(f"ورق سپری: ضخامت {plate_thickness}mm، عرض {plate_width:.0f}mm")
    seismic_checks.append(f"  تسلیم: نسبت = {yield_ratio:.2f} {'✓' if yield_ratio <= 1.0 else '✗'}")
    seismic_checks.append(f"  گسیختگی: نسبت = {rupture_ratio:.2f} {'✓' if rupture_ratio <= 1.0 else '✗'}")
    seismic_checks.append(f"  بلوک برشی: نسبت = {block_shear_ratio:.2f} {'✓' if block_shear_ratio <= 1.0 else '✗'}")
    
    # ========================================================================
    # مرحله 5: طراحی جوش
    # ========================================================================
    
    effective_weld_length = 2 * plate_width + 2 * plate_length
    
    if design_method == DesignMethod.LRFD:
        phi_weld = 0.75
        required_capacity = flange_force_n / phi_weld
    else:
        omega_weld = 2.00
        required_capacity = flange_force_n * omega_weld
    
    weld_area_required = required_capacity / (0.6 * electrode_fexx_mpa)
    weld_throat_required = weld_area_required / effective_weld_length
    weld_size_required = weld_throat_required / 0.707
    
    min_weld_size = max(5.0, math.ceil(plate_thickness / 4.0))
    weld_size = max(weld_size_required, min_weld_size)
    
    if weld_size_mm is not None:
        weld_size = max(weld_size, weld_size_mm)
    
    weld_size = math.ceil(weld_size)
    
    actual_weld_capacity = 0.6 * electrode_fexx_mpa * 0.707 * weld_size * effective_weld_length
    if design_method == DesignMethod.LRFD:
        actual_weld_capacity = phi_weld * actual_weld_capacity
    else:
        actual_weld_capacity = actual_weld_capacity / omega_weld
    
    is_weld_adequate = actual_weld_capacity >= flange_force_n
    weld_utilization = flange_force_n / actual_weld_capacity if actual_weld_capacity > 0 else 1.0
    
    top_weld_result = WeldResult(
        weld_size_mm=weld_size,
        weld_length_mm=effective_weld_length,
        required_size_mm=weld_size_required,
        capacity_n=actual_weld_capacity,
        utilization_ratio=weld_utilization,
        is_adequate=is_weld_adequate
    )
    
    bottom_weld_result = WeldResult(
        weld_size_mm=weld_size,
        weld_length_mm=effective_weld_length,
        required_size_mm=weld_size_required,
        capacity_n=actual_weld_capacity,
        utilization_ratio=weld_utilization,
        is_adequate=is_weld_adequate
    ) if connection_type in [ConnectionType.BOTH_FLANGES, ConnectionType.BOTTOM_FLANGE_ONLY] else None
    
    seismic_checks.append(f"جوش: سایز {weld_size}mm، طول {effective_weld_length:.0f}mm")
    seismic_checks.append(f"  ظرفیت جوش = {actual_weld_capacity/1000:.1f} kN (نسبت مصرف = {weld_utilization:.2f})")
    
    # ========================================================================
    # مرحله 6: کنترل اضافه مقاومت لرزه‌ای
    # ========================================================================
    
    connection_capacity = min(
        num_bolts_required * bolt_capacity_tension * beam_depth_mm,
        actual_capacity_yield * beam_depth_mm,
        actual_weld_capacity * beam_depth_mm
    )
    
    required_capacity = 1.2 * M_pr
    overstrength_ratio = connection_capacity / required_capacity if required_capacity > 0 else 0
    
    if overstrength_ratio >= 1.0:
        seismic_checks.append(f"✓ کنترل اضافه مقاومت: ظرفیت اتصال ≥ 1.2×M_pr ({connection_capacity/1e6:.1f} ≥ {required_capacity/1e6:.1f} kN·m)")
    else:
        seismic_checks.append(f"⚠️ کنترل اضافه مقاومت رد شد: ظرفیت اتصال ({connection_capacity/1e6:.1f}) < 1.2×M_pr ({required_capacity/1e6:.1f} kN·m)")
        seismic_check_passed = False
    
    # ========================================================================
    # جمع‌بندی نهایی
    # ========================================================================
    
    is_valid = (is_bolt_adequate and is_plate_adequate and is_weld_adequate and seismic_check_passed)
    
    summary_parts = []
    if is_valid:
        summary_parts.append("✅ طراحی اتصال BFP با موفقیت انجام شد")
    else:
        summary_parts.append("❌ طراحی اتصال BFP نیاز به اصلاح دارد")
    
    summary_parts.append(f"نوع اتصال: {connection_type.value}")
    summary_parts.append(f"تعداد پیچ: {num_bolts_required} عدد {bolt_type.value} قطر {bolt_diameter_mm}mm")
    summary_parts.append(f"ورق سپری: ضخامت {plate_thickness}mm, عرض {plate_width:.0f}mm")
    summary_parts.append(f"جوش: سایز {weld_size}mm")
    summary_parts.append(f"M_pr = {M_pr/1e6:.1f} kN·m")
    
    if not is_bolt_adequate:
        summary_parts.append("⚠️ پیچ‌ها ظرفیت کافی ندارند")
    if not is_plate_adequate:
        summary_parts.append("⚠️ ورق سپری ظرفیت کافی ندارد")
    if not is_weld_adequate:
        summary_parts.append("⚠️ جوش ظرفیت کافی ندارد")
    
    return BFPConnectionDesign(
        connection_type=connection_type,
        design_method=design_method,
        fy_beam_mpa=fy_beam_mpa,
        fu_beam_mpa=fu_beam_mpa,
        zx_beam_mm3=zx_beam_mm3,
        beam_depth_mm=beam_depth_mm,
        beam_flange_width_mm=beam_flange_width_mm,
        beam_flange_thickness_mm=beam_flange_thickness_mm,
        plate_fy_mpa=plate_fy_mpa,
        plate_fu_mpa=plate_fu_mpa,
        m_pr_nmm=M_pr,
        flange_force_n=flange_force_n,
        top_bolts=top_bolts_result,
        bottom_bolts=bottom_bolts_result,
        top_plate=top_plate_result,
        bottom_plate=bottom_plate_result,
        top_weld=top_weld_result,
        bottom_weld=bottom_weld_result,
        seismic_check_passed=seismic_check_passed,
        seismic_checks=seismic_checks,
        overstrength_ratio=overstrength_ratio,
        is_valid=is_valid,
        summary="\n".join(summary_parts)
    )


# ============================================================================
# توابع خروجی به فرمت Word (DOCX)
# ============================================================================

def _set_cell_font(cell, text: str, font_size_pt: int = 10, bold: bool = False, 
                    alignment: str = "left", font_name: str = "B Nazanin"):
    """تنظیم فونت و متن سلول در جدول Word"""
    cell.text = text
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if alignment == "left" else \
                          WD_ALIGN_PARAGRAPH.CENTER if alignment == "center" else \
                          WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.name = font_name
    
    # پشتیبانی از فارسی
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except:
        pass


def _add_heading(doc, text: str, level: int = 1):
    """اضافه کردن عنوان به سند Word"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.name = "B Nazanin"
        run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    return heading


def _add_paragraph(doc, text: str, font_size_pt: int = 11, bold: bool = False):
    """اضافه کردن پاراگراف به سند Word"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.name = "B Nazanin"
    return p


def export_bfp_to_word(result: BFPConnectionDesign, filename: str = "BFP_Connection_Design.docx") -> str:
    """
    خروجی نتایج طراحی اتصال BFP به فایل Word
    
    Args:
        result: نتایج طراحی از تابع design_bfp_connection
        filename: نام فایل خروجی
    
    Returns:
        مسیر فایل ایجاد شده
    """
    if not DOCX_AVAILABLE:
        raise ImportError("برای خروجی Word به کتابخانه python-docx نیاز است. نصب: pip install python-docx")
    
    doc = Document()
    
    # تنظیم حاشیه صفحات
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ========================================================================
    # عنوان اصلی
    # ========================================================================
    title = doc.add_heading("گزارش طراحی اتصال BFP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "B Nazanin"
        run.font.size = Pt(18)
        run.font.bold = True
    
    # هدر اطلاعات
    info_table = doc.add_table(rows=2, cols=3)
    info_table.style = 'Table Grid'
    
    _set_cell_font(info_table.cell(0, 0), "نوع اتصال:", 11, True)
    _set_cell_font(info_table.cell(0, 1), result.connection_type.value, 11, False)
    _set_cell_font(info_table.cell(0, 2), "", 11, False)
    
    _set_cell_font(info_table.cell(1, 0), "روش طراحی:", 11, True)
    _set_cell_font(info_table.cell(1, 1), result.design_method.value, 11, False)
    _set_cell_font(info_table.cell(1, 2), f"تاریخ طراحی: {result.design_date}", 11, False)
    
    doc.add_paragraph()
    
    # ========================================================================
    # بخش 1: مشخصات تیر
    # ========================================================================
    _add_heading(doc, "1. مشخصات تیر", 2)
    
    beam_data = [
        ("تنش تسلیم تیر (F_y)", f"{result.fy_beam_mpa} MPa"),
        ("تنش نهایی تیر (F_u)", f"{result.fu_beam_mpa} MPa"),
        ("اساس پلاستیک تیر (Z_x)", f"{result.zx_beam_mm3:,.0f} mm³"),
        ("عمق تیر (d)", f"{result.beam_depth_mm} mm"),
        ("عرض بال تیر (b_f)", f"{result.beam_flange_width_mm} mm"),
        ("ضخامت بال تیر (t_f)", f"{result.beam_flange_thickness_mm} mm"),
    ]

    beam_table = doc.add_table(rows=len(beam_data), cols=2)
    beam_table.style = 'Table Grid'
    
    for i, (key, value) in enumerate(beam_data):
        _set_cell_font(beam_table.cell(i, 0), key, 10, True)
        _set_cell_font(beam_table.cell(i, 1), value, 10, False)
    
    doc.add_paragraph()
    
    # ========================================================================
    # بخش 2: نتایج محاسبات اصلی
    # ========================================================================
    _add_heading(doc, "2. نتایج محاسبات اصلی", 2)
    
    main_table = doc.add_table(rows=3, cols=2)
    main_table.style = 'Table Grid'
    
    main_data = [
        ("لنگر پلاستیک مورد انتظار (M_pr)", f"{result.m_pr_nmm/1e6:.1f} kN·m"),
        ("نیروی کششی در ورق سپری (F)", f"{result.flange_force_n/1000:.1f} kN"),
        ("نسبت اضافه مقاومت (اتصال / 1.2×M_pr)", f"{result.overstrength_ratio:.2f}"),
    ]
    
    for i, (key, value) in enumerate(main_data):
        _set_cell_font(main_table.cell(i, 0), key, 10, True)
        _set_cell_font(main_table.cell(i, 1), value, 10, False)
    
    doc.add_paragraph()
    
    # ========================================================================
    # بخش 3: نتایج طراحی پیچ‌ها
    # ========================================================================
    _add_heading(doc, "3. نتایج طراحی پیچ‌ها", 2)
    
    bolt_cols = ["جزء", "تعداد", "قطر (mm)", "نوع", "نیروی هر پیچ (kN)", "ظرفیت (kN)", "نسبت مصرف", "وضعیت"]
    bolt_table = doc.add_table(rows=1, cols=len(bolt_cols))
    bolt_table.style = 'Table Grid'
    
    # هدر جدول
    for i, col in enumerate(bolt_cols):
        _set_cell_font(bolt_table.cell(0, i), col, 10, True, "center")
    
    # داده‌های بالایی
    if result.top_bolts:
        b = result.top_bolts
        row = bolt_table.add_row()
        _set_cell_font(row.cells[0], "ورق بالایی", 10, False, "center")
        _set_cell_font(row.cells[1], str(b.num_bolts), 10, False, "center")
        _set_cell_font(row.cells[2], f"{b.bolt_diameter_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[3], b.bolt_type.value, 10, False, "center")
        _set_cell_font(row.cells[4], f"{b.bolt_force_n/1000:.1f}", 10, False, "center")
        _set_cell_font(row.cells[5], f"{b.bolt_capacity_n/1000:.1f}", 10, False, "center")
        _set_cell_font(row.cells[6], f"{b.utilization_ratio:.2f}", 10, False, "center")
        status = "✓" if b.is_adequate else "✗"
        _set_cell_font(row.cells[7], status, 10, False, "center")
    
    # داده‌های پایینی
    if result.bottom_bolts:
        b = result.bottom_bolts
        row = bolt_table.add_row()
        _set_cell_font(row.cells[0], "ورق پایینی", 10, False, "center")
        _set_cell_font(row.cells[1], str(b.num_bolts), 10, False, "center")
        _set_cell_font(row.cells[2], f"{b.bolt_diameter_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[3], b.bolt_type.value, 10, False, "center")
        _set_cell_font(row.cells[4], f"{b.bolt_force_n/1000:.1f}", 10, False, "center")
        _set_cell_font(row.cells[5], f"{b.bolt_capacity_n/1000:.1f}", 10, False, "center")
        _set_cell_font(row.cells[6], f"{b.utilization_ratio:.2f}", 10, False, "center")
        status = "✓" if b.is_adequate else "✗"
        _set_cell_font(row.cells[7], status, 10, False, "center")
    
    doc.add_paragraph()
    
    # ========================================================================
    # بخش 4: نتایج طراحی ورق سپری
    # ========================================================================
    _add_heading(doc, "4. نتایج طراحی ورق سپری", 2)
    
    plate_cols = ["جزء", "ضخامت (mm)", "عرض (mm)", "طول (mm)", "تسلیم", "گسیختگی", "بلوک برشی", "وضعیت"]
    plate_table = doc.add_table(rows=1, cols=len(plate_cols))
    plate_table.style = 'Table Grid'
    
    for i, col in enumerate(plate_cols):
        _set_cell_font(plate_table.cell(0, i), col, 10, True, "center")
    
    if result.top_plate:
        p = result.top_plate
        row = plate_table.add_row()
        _set_cell_font(row.cells[0], "ورق بالایی", 10, False, "center")
        _set_cell_font(row.cells[1], f"{p.plate_thickness_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[2], f"{p.plate_width_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[3], f"{p.plate_length_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[4], f"{p.yield_check_ratio:.2f}", 10, False, "center")
        _set_cell_font(row.cells[5], f"{p.rupture_check_ratio:.2f}", 10, False, "center")
        _set_cell_font(row.cells[6], f"{p.block_shear_ratio:.2f}", 10, False, "center")
        status = "✓" if p.is_adequate else "✗"
        _set_cell_font(row.cells[7], status, 10, False, "center")
    
    if result.bottom_plate:
        p = result.bottom_plate
        row = plate_table.add_row()
        _set_cell_font(row.cells[0], "ورق پایینی", 10, False, "center")
        _set_cell_font(row.cells[1], f"{p.plate_thickness_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[2], f"{p.plate_width_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[3], f"{p.plate_length_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[4], f"{p.yield_check_ratio:.2f}", 10, False, "center")
        _set_cell_font(row.cells[5], f"{p.rupture_check_ratio:.2f}", 10, False, "center")
        _set_cell_font(row.cells[6], f"{p.block_shear_ratio:.2f}", 10, False, "center")
        status = "✓" if p.is_adequate else "✗"
        _set_cell_font(row.cells[7], status, 10, False, "center")
    
    doc.add_paragraph()
    
    # ========================================================================
    # بخش 5: نتایج طراحی جوش
    # ========================================================================
    _add_heading(doc, "5. نتایج طراحی جوش", 2)
    
    weld_cols = ["جزء", "سایز جوش (mm)", "طول جوش (mm)", "ظرفیت (kN)", "نسبت مصرف", "وضعیت"]
    weld_table = doc.add_table(rows=1, cols=len(weld_cols))
    weld_table.style = 'Table Grid'
    
    for i, col in enumerate(weld_cols):
        _set_cell_font(weld_table.cell(0, i), col, 10, True, "center")
    
    if result.top_weld:
        w = result.top_weld
        row = weld_table.add_row()
        _set_cell_font(row.cells[0], "جوش بالایی", 10, False, "center")
        _set_cell_font(row.cells[1], f"{w.weld_size_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[2], f"{w.weld_length_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[3], f"{w.capacity_n/1000:.1f}", 10, False, "center")
        _set_cell_font(row.cells[4], f"{w.utilization_ratio:.2f}", 10, False, "center")
        status = "✓" if w.is_adequate else "✗"
        _set_cell_font(row.cells[5], status, 10, False, "center")
    
    if result.bottom_weld:
        w = result.bottom_weld
        row = weld_table.add_row()
        _set_cell_font(row.cells[0], "جوش پایینی", 10, False, "center")
        _set_cell_font(row.cells[1], f"{w.weld_size_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[2], f"{w.weld_length_mm:.0f}", 10, False, "center")
        _set_cell_font(row.cells[3], f"{w.capacity_n/1000:.1f}", 10, False, "center")
        _set_cell_font(row.cells[4], f"{w.utilization_ratio:.2f}", 10, False, "center")
        status = "✓" if w.is_adequate else "✗"
        _set_cell_font(row.cells[5], status, 10, False, "center")
    
    doc.add_paragraph()
    
    # ========================================================================
    # بخش 6: کنترل‌های لرزه‌ای
    # ========================================================================
    _add_heading(doc, "6. کنترل‌های لرزه‌ای", 2)
    
    for check in result.seismic_checks:
        _add_paragraph(doc, f"• {check}", 10, False)
    
    seismic_status = "✓ تایید شد" if result.seismic_check_passed else "✗ تایید نشد"
    _add_paragraph(doc, f"\nوضعیت نهایی کنترل لرزه‌ای: {seismic_status}", 11, True)
    
    doc.add_paragraph()
    
    # ========================================================================
    # بخش 7: جمع‌بندی و نتیجه‌گیری
    # ========================================================================
    _add_heading(doc, "7. جمع‌بندی و نتیجه‌گیری", 2)
    
    for line in result.summary.split("\n"):
        _add_paragraph(doc, line, 11, line.startswith("✅") or line.startswith("❌"))
    
    doc.add_paragraph()
    
    # ========================================================================
    # پانویس
    # ========================================================================
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("گزارش طراحی اتصال BFP - بر اساس مبحث دهم مقررات ملی ساختمان (ویرایش ۱۴۰۱)")
    footer_run.font.size = Pt(9)
    footer_run.font.name = "B Nazanin"
    
    # ذخیره فایل
    doc.save(filename)
    return os.path.abspath(filename)


# ============================================================================
# توابع نمایش نتایج در کنسول
# ============================================================================

def print_bfp_design_result(result: BFPConnectionDesign):
    """چاپ نتایج طراحی اتصال BFP به صورت خوانا در کنسول"""
    
    print("\n" + "=" * 80)
    print("📐 نتایج طراحی اتصال BFP (Bolted Flange Plate)")
    print("=" * 80)
    
    print(f"\n📋 روش طراحی: {result.design_method.value}")
    print(f"🔗 نوع اتصال: {result.connection_type.value}")
    
    print("\n" + "-" * 50)
    print("🔩 نتایج طراحی پیچ‌ها")
    print("-" * 50)
    
    if result.top_bolts:
        b = result.top_bolts
        print(f"ورق بالایی: {b.num_bolts} عدد پیچ {b.bolt_type.value} قطر {b.bolt_diameter_mm:.0f}mm")
        print(f"  نیروی هر پیچ = {b.bolt_force_n/1000:.1f} kN")
        print(f"  ظرفیت هر پیچ = {b.bolt_capacity_n/1000:.1f} kN")
        print(f"  نسبت مصرف = {b.utilization_ratio:.2f} {'✓' if b.is_adequate else '✗'}")
    
    if result.bottom_bolts:
        b = result.bottom_bolts
        print(f"ورق پایینی: {b.num_bolts} عدد پیچ {b.bolt_type.value} قطر {b.bolt_diameter_mm:.0f}mm")
        print(f"  نیروی هر پیچ = {b.bolt_force_n/1000:.1f} kN")
        print(f"  ظرفیت هر پیچ = {b.bolt_capacity_n/1000:.1f} kN")
        print(f"  نسبت مصرف = {b.utilization_ratio:.2f} {'✓' if b.is_adequate else '✗'}")
    
    print("\n" + "-" * 50)
    print("📄 نتایج طراحی ورق سپری")
    print("-" * 50)
    
    if result.top_plate:
        p = result.top_plate
        print(f"ورق بالایی: ضخامت {p.plate_thickness_mm:.0f}mm, عرض {p.plate_width_mm:.0f}mm, طول {p.plate_length_mm:.0f}mm")
        print(f"  ضخامت موردنیاز = {p.required_thickness_mm:.1f}mm")
        print(f"  کنترل تسلیم = {p.yield_check_ratio:.2f} {'✓' if p.yield_check_ratio <= 1.0 else '✗'}")
        print(f"  کنترل گسیختگی = {p.rupture_check_ratio:.2f} {'✓' if p.rupture_check_ratio <= 1.0 else '✗'}")
        print(f"  کنترل بلوک برشی = {p.block_shear_ratio:.2f} {'✓' if p.block_shear_ratio <= 1.0 else '✗'}")
    
    if result.bottom_plate:
        p = result.bottom_plate
        print(f"ورق پایینی: ضخامت {p.plate_thickness_mm:.0f}mm, عرض {p.plate_width_mm:.0f}mm, طول {p.plate_length_mm:.0f}mm")
    
    print("\n" + "-" * 50)
    print("🔥 نتایج طراحی جوش")
    print("-" * 50)
    
    if result.top_weld:
        w = result.top_weld
        print(f"جوش بالایی: سایز {w.weld_size_mm:.0f}mm, طول {w.weld_length_mm:.0f}mm")
        print(f"  سایز موردنیاز = {w.required_size_mm:.1f}mm")
        print(f"  ظرفیت = {w.capacity_n/1000:.1f} kN (نسبت مصرف = {w.utilization_ratio:.2f}) {'✓' if w.is_adequate else '✗'}")
    
    if result.bottom_weld:
        w = result.bottom_weld
        print(f"جوش پایینی: سایز {w.weld_size_mm:.0f}mm, طول {w.weld_length_mm:.0f}mm")
    
    print("\n" + "-" * 50)
    print("🌍 کنترل‌های لرزه‌ای")
    print("-" * 50)
    
    for check in result.seismic_checks:
        print(f"  {check}")
    
    print("\n" + "-" * 50)
    print("📊 جمع‌بندی نهایی")
    print("-" * 50)
    print(result.summary)
    
    print("\n" + "=" * 80)
    print(f"✅ وضعیت نهایی: {'طراحی معتبر است' if result.is_valid else 'طراحی نیاز به اصلاح دارد'}")
    
    if DOCX_AVAILABLE:
        print(f"\n📄 برای خروجی Word، تابع export_bfp_to_word(result, 'filename.docx') را فراخوانی کنید.")
    else:
        print("\n⚠️ برای خروجی Word، کتابخانه python-docx را نصب کنید: pip install python-docx")
    print("=" * 80)


# ============================================================================
# مثال استفاده
# ============================================================================

if __name__ == "__main__":
    print("=" * 80)
    print("مثال طراحی اتصال BFP برای قاب خمشی ویژه")
    print("=" * 80)
    
    # مشخصات تیر (مثال: تیر IPE 330)
    fy_beam = 345.0      # MPa - فولاد St-52
    fu_beam = 450.0      # MPa
    zx_beam = 1.02e6     # mm³
    beam_depth = 330.0   # mm
    flange_width = 160.0  # mm
    flange_thickness = 11.5  # mm
    
    # انجام طراحی
    result = design_bfp_connection(
        fy_beam_mpa=fy_beam,
        fu_beam_mpa=fu_beam,
        zx_beam_mm3=zx_beam,
        beam_depth_mm=beam_depth,
        beam_flange_width_mm=flange_width,
        beam_flange_thickness_mm=flange_thickness,
        bolt_type=BoltType.A325,
        bolt_diameter_mm=20.0,
        bolt_grade_ry=1.1,
        plate_fy_mpa=345.0,
        plate_fu_mpa=450.0,
        edge_distance_mm=40.0,
        bolt_spacing_mm=80.0,
        electrode_fexx_mpa=480.0,
        design_method=DesignMethod.LRFD,
        connection_type=ConnectionType.BOTH_FLANGES,
        max_num_bolts=12,
        min_plate_thickness_mm=12.0
    )
    
    # نمایش نتایج در کنسول
    print_bfp_design_result(result)
    
    # خروجی به Word (در صورت نصب کتابخانه)
    if DOCX_AVAILABLE:
        try:
            filepath = export_bfp_to_word(result, r"E:\BFP_Connection_Design_Report.docx")
            print(f"\n📄 فایل Word با موفقیت ذخیره شد: {filepath}")
        except Exception as e:
            print(f"\n⚠️ خطا در ذخیره فایل Word: {e}")
    else:
        print("\n⚠️ برای خروجی Word، ابتدا کتابخانه را نصب کنید: pip install python-docx")