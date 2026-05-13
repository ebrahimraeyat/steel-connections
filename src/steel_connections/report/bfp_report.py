# -*- coding: utf-8 -*-
"""
Word (.docx) design report generator for BFP (Bolted Flange Plate) connections.

Usage
-----
    from steel_connections.report.bfp_report import generate_report
    path = generate_report(connection, project_info={...})
"""

from __future__ import annotations

import io
import math
import tempfile
from datetime import date
from pathlib import Path
from typing import TYPE_CHECKING

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if TYPE_CHECKING:
    from steel_connections.bfp_connection import BFPConnection


# ── OMML equation helper ──────────────────────────────────────────────────────
# Inserts a Word OMML math block (inline or display) into a paragraph.
# We build the XML manually for common formula patterns.

_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def _omml(xml_inner: str) -> etree._Element:
    """Wrap OMML inner content in an <m:oMath> block."""
    xml = (
        f'<m:oMath xmlns:m="{_OMML_NS}">'
        f'{xml_inner}'
        f'</m:oMath>'
    )
    return etree.fromstring(xml)


def _mrun(text: str, italic: bool = False, bold: bool = False) -> str:
    style = ""
    if italic:
        style += "<m:i/>"
    if bold:
        style += "<m:b/>"
    rpr = f"<m:rPr>{style}</m:rPr>" if style else ""
    return f"<m:r>{rpr}<m:t>{text}</m:t></m:r>"


def _mfrac(num: str, den: str) -> str:
    return f"<m:f><m:num><m:r><m:t>{num}</m:t></m:r></m:num><m:den><m:r><m:t>{den}</m:t></m:r></m:den></m:f>"


def _eq_paragraph(doc: Document, omml_xml: str, center: bool = True) -> None:
    """Add a display equation paragraph."""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_el = _omml(omml_xml)
    p._p.append(math_el)


# ── Style helpers ─────────────────────────────────────────────────────────────

_ACCENT = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
_OK     = RGBColor(0x29, 0x7A, 0x35)   # green
_FAIL   = RGBColor(0xC0, 0x00, 0x00)   # red
_HEAD   = RGBColor(0xFF, 0xFF, 0xFF)   # white (for table headers)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _header_row(table, col_labels: list[str], bg: str = "1F497D") -> None:
    row = table.rows[0]
    for i, label in enumerate(col_labels):
        cell = row.cells[i]
        _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold   = True
        run.font.color.rgb = _HEAD
        run.font.size = Pt(9)


def _make_table(doc: Document, headers: list[str],
                rows: list[list[str]], style: str = "Table Grid") -> None:
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(table, headers)
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()   # spacing


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = "Normal"
    for run in p.runs:
        run.font.size = Pt(10)


def _formula_line(doc: Document, label: str, formula: str,
                  result: str, unit: str = "") -> None:
    """Single-line formula display: label   formula = result unit"""
    p = doc.add_paragraph(style="Normal")
    r = p.add_run(f"    {label}  →  {formula} = ")
    r.font.size  = Pt(10)
    r.font.italic = True
    rv = p.add_run(f"{result}  {unit}")
    rv.font.size  = Pt(10)
    rv.bold       = True


def _check_line(doc: Document, label: str, demand: str, capacity: str,
                ok: bool, dcr: float | None = None) -> None:
    p = doc.add_paragraph(style="Normal")
    sym = "✓" if ok else "✗"
    dcr_str = f"  (DCR = {dcr:.2f})" if dcr is not None else ""
    r = p.add_run(f"    {sym}  {label}:  {demand}  <  {capacity}{dcr_str}  →  {'OK' if ok else 'FAIL'}")
    r.font.size  = Pt(10)
    r.font.bold  = True
    r.font.color.rgb = _OK if ok else _FAIL


def _fmt(v, digits: int = 3) -> str:
    try:
        if math.isnan(v):
            return "—"
        return f"{v:.{digits}f}"
    except Exception:
        return str(v)


def _section_break(doc: Document) -> None:
    doc.add_paragraph()


# ── Main report function ──────────────────────────────────────────────────────

def generate_report(connection: "BFPConnection",
                    project_info: dict | None = None,
                    output_path: str | Path | None = None,
                    view_images: dict | None = None) -> Path:
    """
    Generate a Word design report for a BFP connection.

    Parameters
    ----------
    connection   : BFPConnection instance (fully constructed).
    project_info : dict with keys: project, engineer, checker, firm,
                   member, level, date (all optional).
    output_path  : Where to save the .docx. Defaults to current directory.

    Returns
    -------
    Path to the generated .docx file.
    """
    pi = project_info or {}
    doc = Document()

    # ── page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER HEADER ─────────────────────────────────────────────────────────
    design_code_name = getattr(connection, 'design_code', '—')
    _heading(doc, "Bolted Flange Plate (BFP) Connection — Design Report", level=1)

    hdr_rows = [
        ["Project",  pi.get("project",  "—"), "Engineer", pi.get("engineer", "—")],
        ["Member",   pi.get("member",   "—"), "Checker",  pi.get("checker",  "—")],
        ["Level",    pi.get("level",    "—"), "Firm",     pi.get("firm",     "—")],
        ["Date",     pi.get("date",     str(date.today())), "Page", "1"],
        ["Design Code", design_code_name, "", ""],
    ]
    tbl = doc.add_table(rows=len(hdr_rows), cols=4)
    tbl.style = "Table Grid"
    for r_i, row_data in enumerate(hdr_rows):
        row = tbl.rows[r_i]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_i % 2 == 0:
                run.bold = True
                _set_cell_bg(cell, "D9E1F2")
    doc.add_paragraph()

    # ── 3D CONNECTION VIEWS ───────────────────────────────────────────────────
    if view_images:
        _heading(doc, "3D Connection Views", level=2)
        _view_labels = {
            "iso":   "Isometric View",
            "front": "Front View (looking along beam)",
            "side":  "Side View (beam web)",
            "top":   "Top View (flange)",
        }
        # Build a 2 × N table for pairs of images
        items = [(k, view_images[k]) for k in ("iso", "front", "side", "top") if k in view_images]
        # 2 images per row
        row_count = (len(items) + 1) // 2
        if items:
            vtbl = doc.add_table(rows=row_count * 2, cols=2)  # label row + image row per pair
            vtbl.style = "Table Grid"
            vtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (key, img_path) in enumerate(items):
                row_pair = i // 2
                col_idx  = i % 2
                # label row
                lbl_cell = vtbl.rows[row_pair * 2].cells[col_idx]
                lp = lbl_cell.paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lr = lp.add_run(_view_labels.get(key, key))
                lr.bold = True
                lr.font.size = Pt(9)
                # image row
                img_cell = vtbl.rows[row_pair * 2 + 1].cells[col_idx]
                ip = img_cell.paragraphs[0]
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    run = ip.add_run()
                    run.add_picture(img_path, width=Inches(3.0))
                except Exception:
                    ip.add_run(f"[{key} image unavailable]")
        doc.add_paragraph()

    # ── 1. MEMBER PROPERTIES ─────────────────────────────────────────────────
    _heading(doc, "1  Member Properties", level=2)

    beam = connection.beam
    col  = connection.column
    bg   = beam.geom
    cg   = col.geom

    # Beam
    _heading(doc, "1.1  Beam Section", level=3)
    _make_table(doc,
        headers=["Property", "Symbol", "Value", "Unit"],
        rows=[
            ["Total depth",          "d",   _fmt(bg.d),   "cm"],
            ["Flange width",         "bf",  _fmt(bg.b),   "cm"],
            ["Flange thickness",     "tf",  _fmt(bg.t_f), "cm"],
            ["Web thickness",        "tw",  _fmt(bg.t_w), "cm"],
            ["Plastic modulus (X)",  "Zx",  _fmt(bg.Z_x), "cm³"],
            ["Elastic modulus (X)",  "Sx",  _fmt(bg.S_x), "cm³"],
            ["Cross-section area",   "Ag",  _fmt(bg.A_g), "cm²"],
            ["Yield strength",       "fy",  _fmt(beam.mat.f_y), "kg/cm²"],
            ["Tensile strength",     "fu",  _fmt(beam.mat.f_u), "kg/cm²"],
        ]
    )

    # Column
    _heading(doc, "1.2  Column Section", level=3)
    _make_table(doc,
        headers=["Property", "Symbol", "Value", "Unit"],
        rows=[
            ["Total depth",      "d",  _fmt(cg.d),   "cm"],
            ["Flange width",     "bf", _fmt(cg.b),   "cm"],
            ["Flange thickness", "tf", _fmt(cg.t_f), "cm"],
            ["Web thickness",    "tw", _fmt(cg.t_w), "cm"],
        ]
    )

    # ── 2. CONNECTION COMPONENTS ──────────────────────────────────────────────
    _heading(doc, "2  Connection Components", level=2)

    plate = connection.plate
    bolt  = connection.bolt
    bg2   = connection.bolt_group

    _heading(doc, "2.1  Flange Plate", level=3)
    _make_table(doc,
        headers=["Property", "Symbol", "Value", "Unit"],
        rows=[
            ["Plate width",       "bp",  _fmt(plate.b_i),  "cm"],
            ["Plate length",      "Lp",  _fmt(plate.h_i),  "cm"],
            ["Plate thickness",   "tp",  _fmt(plate.t_i),  "cm"],
            ["Yield strength",    "Fyp", _fmt(plate.f_yi), "kg/cm²"],
            ["Tensile strength",  "Fup", _fmt(plate.f_ui), "kg/cm²"],
        ]
    )

    _heading(doc, "2.2  Bolt Group (Flange)", level=3)
    _make_table(doc,
        headers=["Property", "Symbol", "Value", "Unit"],
        rows=[
            ["Bolt diameter",         "df",  _fmt(bolt.d_f),       "cm"],
            ["Bolt tensile strength", "Fuf", _fmt(bolt.f_uf),      "kg/cm²"],
            ["Bolt shear area",       "Ao",  _fmt(bolt.A_o),        "cm²"],
            ["Bolt rows (n_p)",       "np",  str(bg2.n_p),          "—"],
            ["Bolt gauge lines (n_g)", "ng", str(bg2.n_g),          "—"],
            ["Row spacing (s_p)",     "sp",  _fmt(bg2.s_p),         "cm"],
            ["Gauge spacing (s_g)",   "sg",  _fmt(bg2.s_g),         "cm"],
            ["Total bolts / plate",   "nb",  str(bg2.n_b),          "—"],
        ]
    )

    # ── 3. DESIGN PROCEDURE ───────────────────────────────────────────────────
    cr = getattr(connection, 'code_refs', {})
    _heading(doc, "3  Design Procedure", level=2)
    _body(doc,
        f"The BFP connection is designed following a capacity-based approach per "
        f"{design_code_name}. "
        "The probable maximum moment at the plastic hinge is used to determine "
        "required plate forces and bolt demands.")

    # ── Step 1 ─────────────────────────────────────
    _heading(doc, f"Step 1 — Probable Maximum Moment at Plastic Hinge  [{cr.get('mpr','')  }]", level=3)

    cpr = connection.cpr
    mp  = connection.m_p
    mpr = connection.m_pr
    fy  = beam.mat.f_y
    fu  = beam.mat.f_u
    Ry  = beam.Ry
    Zx  = bg.Z_x

    _body(doc, f"[{cr.get('cpr','')}]  Strain hardening factor C_pr:")
    _formula_line(doc, "C_pr",
        f"min( (fy + fu) / (2·fy), 1.2 )  =  min( ({_fmt(fy,0)} + {_fmt(fu,0)}) / (2×{_fmt(fy,0)}), 1.2 )",
        _fmt(cpr, 3))

    _body(doc, f"[{cr.get('mp','')}]  Plastic moment of the beam:")
    _formula_line(doc, "M_p",
        f"Zx × fy  =  {_fmt(Zx)} cm³ × {_fmt(fy,0)} kg/cm²",
        _fmt(mp, 1), "kg·cm")

    _body(doc, f"[{cr.get('mpr','')}]  Probable maximum moment at the plastic hinge:")
    _formula_line(doc, "M_pr",
        f"C_pr × Ry × M_p  =  {_fmt(cpr)} × {_fmt(Ry)} × {_fmt(mp,1)}",
        _fmt(mpr, 1), "kg·cm")

    # ── Step 2 ─────────────────────────────────────
    _heading(doc, f"Step 2 — Maximum Bolt Diameter  [{cr.get('bolt_diam','')}]", level=3)
    db_max = connection.get_max_bolt_diameter()
    _body(doc, f"[{cr.get('bolt_diam','')}]  Max bolt diameter to avoid net-section failure of beam flange:")
    _formula_line(doc, "d_b,max",
        f"bf/2 × (1 − Ry·fy/(Rt·fu)) − offset",
        _fmt(db_max, 3), "cm")
    ok_diam = bolt.d_f <= db_max
    _check_line(doc, "Bolt diameter check",
                f"df = {_fmt(bolt.d_f)} cm", f"db,max = {_fmt(db_max,3)} cm", ok_diam)

    # ── Step 3 ─────────────────────────────────────
    _heading(doc, f"Step 3 — Nominal Shear Force per Bolt  [{cr.get('bolt_shear','')}]", level=3)
    rn1, rn2, rn3 = connection.nominal_shear_force_of_bolt_values()
    rn = min(rn1, rn2, rn3)
    _body(doc, f"[{cr.get('bolt_shear','')}]  Design shear per bolt — minimum of three limit states:")
    _formula_line(doc, "φRn,1 (bolt shear)",
        f"nominal bolt shear capacity",
        _fmt(rn1, 2), "t")
    _formula_line(doc, "φRn,2 (beam flange bearing)",
        f"2.4 × Fuf × df × tf  =  2.4 × {_fmt(bolt.f_uf,0)} × {_fmt(bolt.d_f)} × {_fmt(bg.t_f)}",
        _fmt(rn2, 2), "kg")
    _formula_line(doc, "φRn,3 (plate bearing)",
        f"2.4 × Fup × df × tp  =  2.4 × {_fmt(plate.f_ui,0)} × {_fmt(bolt.d_f)} × {_fmt(plate.t_i)}",
        _fmt(rn3, 2), "kg")
    _body(doc, f"    →  φRn (governs) = {_fmt(rn, 2)}  kg")

    # ── Step 4 ─────────────────────────────────────
    _heading(doc, f"Step 4 — Minimum Number of Bolts per Plate  [{cr.get('n_bolts','')}]", level=3)
    n_min = connection.min_no_bolts()
    _body(doc, f"[{cr.get('n_bolts','')}]  Minimum bolt count to resist probable moment:")
    _formula_line(doc, "n_min",
        f"1.25 × M_pr / (φ × φRn × (d + tp))",
        str(n_min), "bolts")
    ok_nb = bg2.n_b >= n_min
    _check_line(doc, "Bolt count check",
                f"n_provided = {bg2.n_b}", f"n_min = {n_min}", ok_nb)

    # ── Step 5 ─────────────────────────────────────
    _heading(doc, f"Step 5 — Connection Geometry (sh, lh)  [{cr.get('sh_lh','')}]", level=3)
    sh = connection.sh
    lh = connection.lh
    _body(doc, "Distance from column face to centre of last bolt row:")
    _formula_line(doc, "sh",
        f"s1 + sp × (np − 1)  =  {_fmt(connection.s1)} + {_fmt(bg2.s_p)} × ({bg2.n_p / 2} − 1)",
        _fmt(sh, 2), "cm")
    _body(doc, "Clear length between plastic hinges:")
    _formula_line(doc, "lh",
        f"L − 2×sh  =  {_fmt(connection.beam_length)} − 2×{_fmt(sh,2)}",
        _fmt(lh, 2), "cm")
    _body(doc, "Edge distances:")
    _formula_line(doc, "s3 (plate edge)",
        f"(bp − sg) / 2  =  ({_fmt(plate.b_i)} − {_fmt(bg2.s_g)}) / 2",
        _fmt(connection.s3, 2), "cm")
    _formula_line(doc, "s5 (flange edge)",
        f"(bf − sg) / 2  =  ({_fmt(bg.b)} − {_fmt(bg2.s_g)}) / 2",
        _fmt(connection.s5, 2), "cm")
    _formula_line(doc, "kl (effective length for buckling)",
        f"0.65 × s1  =  0.65 × {_fmt(connection.s1)}",
        _fmt(connection.kl, 2), "cm")

    # ── Step 6–8 ────────────────────────────────────
    _heading(doc,
        f"Steps 6–8 — Shear at Hinge, M_f, F_pr  "
        f"[{cr.get('vh','')} / {cr.get('mf','')} / {cr.get('fpr','')}]",
        level=3)
    v_assumed = 0.0
    vh = connection.shear_in_hinge(v_assumed)
    mf = connection.probable_moment_in_column_face(v_assumed)
    fpr = connection.force_of_plate(v_assumed)

    _body(doc, f"[{cr.get('vh','')}]  Shear at plastic hinge (V_gravity = 0):")
    _formula_line(doc, "V_h",
        f"2×M_pr / lh + V  =  2×{_fmt(mpr,1)} / {_fmt(lh,2)} + {_fmt(v_assumed,0)}",
        _fmt(vh, 2), "kg")

    _body(doc, f"[{cr.get('mf','')}]  Probable moment at column face:")
    _formula_line(doc, "M_f",
        f"M_pr + V_h × sh  =  {_fmt(mpr,1)} + {_fmt(vh,2)} × {_fmt(sh,2)}",
        _fmt(mf, 1), "kg·cm")

    _body(doc, f"[{cr.get('fpr','')}]  Required plate force:")
    _formula_line(doc, "F_pr",
        f"M_f / (d + tp)  =  {_fmt(mf,1)} / ({_fmt(bg.d)} + {_fmt(plate.t_i)})",
        _fmt(fpr, 1), "kg")

    # ── Step 9 ─────────────────────────────────────
    _heading(doc, f"Step 9 — Required Bolt Count Under F_pr  [{cr.get('n_bolts','')}]", level=3)
    n9 = connection.check_no_of_bolts(v_assumed)
    _body(doc, "Number of bolts required to resist the actual plate force:")
    _formula_line(doc, "n_req",
        f"F_pr / (φ × φRn)  =  {_fmt(fpr,1)} / (0.9 × {_fmt(rn,2)})",
        str(n9), "bolts")
    ok_n9 = bg2.n_b >= n9
    _check_line(doc, "Bolt count vs F_pr",
                f"n = {bg2.n_b}", f"n_req = {n9}", ok_n9)

    # ── Step 10 ────────────────────────────────────
    _heading(doc, f"Step 10 — Minimum Plate Thickness  [{cr.get('t_min','')}]", level=3)
    t_min = connection.get_minimum_thickness_of_plate(v_assumed)
    _body(doc, f"[{cr.get('t_min','')}]  Required plate thickness (yield limit state):")
    _formula_line(doc, "t_min",
        f"F_pr / (φd × Fyp × bp)  =  {_fmt(fpr,1)} / (1.0 × {_fmt(plate.f_yi,0)} × {_fmt(plate.b_i)})",
        _fmt(t_min, 3), "cm")
    ok_t = plate.t_i >= t_min
    _check_line(doc, "Plate thickness check",
                f"tp = {_fmt(plate.t_i)} cm", f"t_min = {_fmt(t_min,3)} cm", ok_t)

    # ── Step 11–12 — Plate rupture / block shear ───
    _heading(doc,
        f"Steps 11–12 — Tensile Rupture and Block Shear  "
        f"[{cr.get('rupture','')} / {cr.get('block_shear','')}]",
        level=3)

    rn_rup  = connection.max_flange_plate_force_according_to_the_limit_state_of_tensile_rupture()
    rn_bs   = connection.flange_plate_force_block_shear()
    phi_n   = 0.9

    _body(doc, f"[{cr.get('rupture','')}]")
    _formula_line(doc, "φRn,rup (tensile rupture)",
        f"Fup × Anp  =  {_fmt(plate.f_ui,0)} × {_fmt(connection.get_net_area_of_plate(),3)}",
        _fmt(rn_rup, 1), "kg")
    ok_rup = (fpr / 2) < phi_n * rn_rup
    _check_line(doc, "Tensile rupture",
                f"Fpr/2 = {_fmt(fpr/2,1)} kg",
                f"φRn = {_fmt(phi_n*rn_rup,1)} kg", ok_rup,
                dcr=(fpr/2) / (phi_n * rn_rup) if rn_rup else None)

    _body(doc, f"[{cr.get('block_shear','')}]")
    _formula_line(doc, "φRn,bs (block shear)",
        f"min(0.6·Fup·Anv + Fup·Ant,  0.6·Fyp·Agv + Fup·Ant)",
        _fmt(rn_bs, 1), "kg")
    ok_bs = connection.check_flange_plate_block_shear(v_assumed)
    _check_line(doc, "Block shear",
                f"Fpr/2 = {_fmt(fpr/2,1)} kg",
                f"φRn,bs = {_fmt(phi_n*rn_bs,1)} kg", ok_bs)

    # ── Step 13 — Compression buckling ────────────
    _heading(doc, f"Step 13 — Plate Compression Buckling  [{cr.get('buckling','')}]", level=3)
    kl_r = connection.buckling_factor_of_plate()
    rn_buck = connection.plate_force_compresion_buckling()
    _body(doc, f"[{cr.get('buckling','')}]")
    _formula_line(doc, "KL/r",
        f"0.65×s1 / rp  =  {_fmt(connection.kl,2)} / {_fmt(plate.r_p,3)}",
        _fmt(kl_r, 2))
    ok_buck = kl_r <= 25
    _check_line(doc, "Buckling check (KL/r ≤ 25)",
                f"KL/r = {_fmt(kl_r,2)}", "25", ok_buck)
    if rn_buck is not None:
        _formula_line(doc, "φRn,buck",
            f"Fyp × Ap  =  {_fmt(plate.f_yi,0)} × {_fmt(plate.A_p,3)}",
            _fmt(rn_buck, 1), "kg")

    # ── 4. DESIGN CHECKS SUMMARY ─────────────────────────────────────────────
    _section_break(doc)
    _heading(doc, "4  Design Checks Summary", level=2)
    _body(doc, f"Design Code: {design_code_name}")

    errors = connection.check_connection()
    err_values = {e.value for e in errors}

    # Build check table — code-agnostic: use error enum's .description and cr dict
    from steel_connections.bfp_connection import BFPCONNECTIONERROR
    from steel_connections.bfp_connection_aisc358 import AISC358BFPConnection, AISC358BFPERROR

    if isinstance(connection, AISC358BFPConnection):
        checks = [
            (AISC358BFPERROR.beam_weight,   "Beam weight",          "≤ 175 lb/ft (≈260 kg/m)", cr.get('preq_beam','')),
            (AISC358BFPERROR.beam_depth,    "Beam depth",           "≤ W36 (≈91.4 cm)",         cr.get('preq_beam','')),
            (AISC358BFPERROR.max_bolt_diameter, "Max bolt diameter",f"≤ {_fmt(db_max,3)} cm",   cr.get('bolt_diam','')),
            (AISC358BFPERROR.minimum_bolt_grade,"Bolt grade",       "Fu ≥ A325 (120 ksi)",      cr.get('preq_bolt','')),
            (AISC358BFPERROR.plate_buckling,"Plate buckling (KL/r)","≤ 25",                     cr.get('buckling','')),
            (AISC358BFPERROR.max_sh,        "sh ≤ d",               f"sh={_fmt(sh,2)} cm",       cr.get('sh_lh','')),
            (AISC358BFPERROR.minimum_s3,    "s3 (plate edge)",      f"≥ 1.5–2×df",              cr.get('sh_lh','')),
            (AISC358BFPERROR.minimum_s5,    "s5 (flange edge)",     f"≥ 1.5–2×df",              cr.get('sh_lh','')),
        ]
    else:
        checks = [
            (BFPCONNECTIONERROR.beam_weight,   "Beam weight",           "≤ 250 kg/m",                cr.get('preq_beam','')),
            (BFPCONNECTIONERROR.beam_depth,    "Beam depth",            "≤ 100 cm",                  cr.get('preq_beam','')),
            (BFPCONNECTIONERROR.max_bolt_diameter,"Max bolt diameter",  "≤ 2.7 cm",                  cr.get('bolt_diam','')),
            (BFPCONNECTIONERROR.minimum_grade_of_bolt,"Bolt grade",     "fuf ≥ 10000 kg/cm²",        cr.get('preq_bolt','')),
            (BFPCONNECTIONERROR.check_max_buckling_factor_of_plate,"Plate buckling (KL/r)","≤ 25",  cr.get('buckling','')),
            (BFPCONNECTIONERROR.max_sh,        "sh ≤ d",                f"sh={_fmt(sh,2)} cm",        cr.get('sh_lh','')),
            (BFPCONNECTIONERROR.minimum_s3,    "s3 (plate edge)",       f"≥ 1.5–2×df",               cr.get('sh_lh','')),
            (BFPCONNECTIONERROR.minimum_s5,    "s5 (flange edge)",      f"≥ 1.5–2×df",               cr.get('sh_lh','')),
        ]

    check_rows = []
    for err_enum, label, criterion, ref in checks:
        passed = err_enum.value not in err_values
        check_rows.append([label, criterion, ref, "✓  OK" if passed else "✗  FAIL"])

    tbl = doc.add_table(rows=1 + len(check_rows), cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(tbl, ["Check", "Criterion", "Code Reference", "Result"])
    for r_i, row_data in enumerate(check_rows):
        row = tbl.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_i == 3:   # Result column
                ok = "OK" in val
                run.font.color.rgb = _OK if ok else _FAIL
                run.bold = True
    doc.add_paragraph()

    # ── overall result line ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not errors:
        run = p.add_run(f"✓  Connection is ADEQUATE — all checks passed per {design_code_name}.")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = _OK
    else:
        run = p.add_run(
            f"✗  Connection has {len(errors)} FAILED check(s) per {design_code_name} — see table above.")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = _FAIL

    # ── save ─────────────────────────────────────────────────────────────────
    if output_path is None:
        output_path = Path.cwd() / "BFP_Connection_Report.docx"
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
