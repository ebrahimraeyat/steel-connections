# -*- coding: utf-8 -*-
import math
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from steel_connections.bfp_connection_design import BFPConnectionDesign, DesignMethod

_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def _omml(xml_inner: str) -> etree._Element:
    xml = f'<m:oMath xmlns:m="{_OMML_NS}">{xml_inner}</m:oMath>'
    return etree.fromstring(xml)

def _mrun(text: str, italic: bool = False, bold: bool = False) -> str:
    style = ""
    if italic: style += "<m:i/>"
    if bold: style += "<m:b/>"
    rpr = f"<m:rPr>{style}</m:rPr>" if style else ""
    return f"<m:r>{rpr}<m:t>{text}</m:t></m:r>"

def _mfrac(num: str, den: str) -> str:
    return f"<m:f><m:num><m:r><m:t>{num}</m:t></m:r></m:num><m:den><m:r><m:t>{den}</m:t></m:r></m:den></m:f>"

def _eq_paragraph(doc: Document, omml_xml: str, center: bool = True) -> None:
    p = doc.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_el = _omml(omml_xml)
    p._p.append(math_el)

_ACCENT = RGBColor(0x1F, 0x49, 0x7D)
_OK     = RGBColor(0x29, 0x7A, 0x35)
_FAIL   = RGBColor(0xC0, 0x00, 0x00)
_HEAD   = RGBColor(0xFF, 0xFF, 0xFF)

def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _header_row(table, col_labels: list[str], bg: str = "1F497D") -> None:
    row = table.rows[0]
    for i, label in enumerate(col_labels):
        cell = row.cells[i]
        _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = _HEAD
        run.font.size = Pt(9)

def _make_table(doc: Document, headers: list[str], rows: list[list[str]], style: str = "Table Grid") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(table, headers)
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()

def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'B Nazanin'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')

def _body(doc: Document, text: str, rtl: bool = False) -> None:
    p = doc.add_paragraph(text)
    p.style = "Normal"
    if rtl:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.bidi = True
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'B Nazanin'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')

def _formula_line(doc: Document, label: str, formula: str, result: str, unit: str = "", rtl: bool = False) -> None:
    p = doc.add_paragraph(style="Normal")
    if rtl:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.bidi = True
        r = p.add_run(f"    {label}  →  {formula} = ")
    else:
        r = p.add_run(f"    {label}  →  {formula} = ")
    r.font.size = Pt(10)
    r.font.italic = True
    rv = p.add_run(f"{result}  {unit}")
    rv.font.size = Pt(10)
    rv.bold = True

def _check_line(doc: Document, label: str, demand: str, capacity: str, ok: bool, dcr: float = None, rtl: bool = False) -> None:
    p = doc.add_paragraph(style="Normal")
    if rtl:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.bidi = True
    sym = "✓" if ok else "✗"
    dcr_str = f"  (DCR = {dcr:.2f})" if dcr is not None else ""
    r = p.add_run(f"    {sym}  {label}:  {demand}  ≤  {capacity}{dcr_str}  →  {'OK' if ok else 'FAIL'}")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = _OK if ok else _FAIL

def _fmt(v, digits: int = 2) -> str:
    try:
        if math.isnan(v): return "—"
        return f"{v:.{digits}f}"
    except Exception: return str(v)

def generate_mabhas10_report(result: BFPConnectionDesign, project_info: dict = None, output_path: str = None, view_images: dict = None, lang: str = "fa") -> Path:
    from datetime import date

    pi = project_info or {}
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    is_fa = (lang == "fa")

    def txt(fa: str, en: str) -> str:
        return fa if is_fa else en

    std_name = pi.get("standard", txt("مبحث دهم / AISC", "Mabhas 10 / AISC"))
    title_str = txt(
        f"گزارش طراحی اتصال فلنجی پیچی ({std_name})",
        f"BFP Connection Design Report ({std_name})",
    )
    _heading(doc, title_str, level=1)

    hdr_rows = [
        [txt("پروژه", "Project"), pi.get("project", "—"), txt("طراح", "Engineer"), pi.get("engineer", "—")],
        [txt("عضو", "Member"), pi.get("member", "—"), txt("کنترل", "Checker"), pi.get("checker", "—")],
        [txt("آیین‌نامه", "Design Code"), std_name, txt("تاریخ", "Date"), pi.get("date", str(date.today()))],
        [txt("روش طراحی", "Design Method"), result.design_method.value, txt("صفحه", "Page"), "1"],
    ]
    tbl = doc.add_table(rows=len(hdr_rows), cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_i, row_data in enumerate(hdr_rows):
        row = tbl.rows[r_i]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            p = cell.paragraphs[0]
            if is_fa:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            if c_i % 2 == 0:
                run.bold = True
                _set_cell_bg(cell, "D9E1F2")
    doc.add_paragraph()

    if view_images:
        _heading(doc, txt("نماهای سه‌بعدی اتصال", "3D Connection Views"), level=2)
        labels = {
            "iso": txt("نمای ایزومتریک", "Isometric View"),
            "front": txt("نمای روبرو", "Front View"),
            "side": txt("نمای جانبی", "Side View"),
            "top": txt("نمای بالا", "Top View"),
        }
        items = [(k, view_images[k]) for k in ("iso", "front", "side", "top") if k in view_images]
        if items:
            vtbl = doc.add_table(rows=((len(items) + 1) // 2) * 2, cols=2)
            vtbl.style = "Table Grid"
            vtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (key, img_path) in enumerate(items):
                row_pair = i // 2
                col_idx = i % 2
                lp = vtbl.rows[row_pair * 2].cells[col_idx].paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lr = lp.add_run(labels.get(key, key))
                lr.bold = True
                lr.font.size = Pt(9)
                ip = vtbl.rows[row_pair * 2 + 1].cells[col_idx].paragraphs[0]
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    ip.add_run().add_picture(img_path, width=Inches(3.0))
                except Exception:
                    ip.add_run(f"[{key} unavailable]")
            doc.add_paragraph()

    b = result.top_bolts
    p = result.top_plate
    w = result.top_weld
    cpr = min((result.fy_beam_mpa + result.fu_beam_mpa) / (2.0 * result.fy_beam_mpa), 1.2)
    zx_cm3 = result.zx_beam_mm3 / 1000.0
    mpr_knm = result.m_pr_nmm / 1e6
    flange_force_kn = result.flange_force_n / 1000.0
    ry = result.m_pr_nmm / (cpr * result.fy_beam_mpa * result.zx_beam_mm3) if cpr * result.fy_beam_mpa * result.zx_beam_mm3 else 0.0

    _heading(doc, txt("1. مشخصات اعضا", "1. Member Properties"), level=2)
    _make_table(doc,
        headers=[txt("پارامتر", "Property"), txt("نماد", "Symbol"), txt("مقدار", "Value"), txt("واحد", "Unit")],
        rows=[
            [txt("عمق تیر", "Beam depth"), "d", _fmt(pi.get("beam_d_cm", result.beam_depth_mm / 10.0)), "cm"],
            [txt("عرض بال تیر", "Beam flange width"), "bf", _fmt(pi.get("beam_bf_cm", result.beam_flange_width_mm / 10.0)), "cm"],
            [txt("ضخامت بال تیر", "Beam flange thickness"), "tf", _fmt(pi.get("beam_tf_cm", result.beam_flange_thickness_mm / 10.0)), "cm"],
            [txt("ضخامت جان تیر", "Beam web thickness"), "tw", _fmt(pi.get("beam_tw_cm", "—")), "cm"],
            [txt("مدول پلاستیک تیر", "Beam plastic modulus"), "Zx", _fmt(zx_cm3), "cm³"],
            [txt("تنش تسلیم تیر", "Beam yield strength"), "Fy", _fmt(result.fy_beam_mpa), "MPa"],
            [txt("تنش نهایی تیر", "Beam tensile strength"), "Fu", _fmt(result.fu_beam_mpa), "MPa"],
        ]
    )
    _make_table(doc,
        headers=[txt("پارامتر", "Property"), txt("نماد", "Symbol"), txt("مقدار", "Value"), txt("واحد", "Unit")],
        rows=[
            [txt("عمق ستون", "Column depth"), "d", _fmt(pi.get("column_d_cm", "—")), "cm"],
            [txt("عرض بال ستون", "Column flange width"), "bf", _fmt(pi.get("column_bf_cm", "—")), "cm"],
            [txt("ضخامت بال ستون", "Column flange thickness"), "tf", _fmt(pi.get("column_tf_cm", "—")), "cm"],
            [txt("ضخامت جان ستون", "Column web thickness"), "tw", _fmt(pi.get("column_tw_cm", "—")), "cm"],
        ]
    )

    _heading(doc, txt("2. مشخصات اتصال", "2. Connection Components"), level=2)
    if p:
        _make_table(doc,
            headers=[txt("پارامتر", "Property"), txt("نماد", "Symbol"), txt("مقدار", "Value"), txt("واحد", "Unit")],
            rows=[
                [txt("عرض ورق", "Plate width"), "bp", _fmt(p.plate_width_mm), "mm"],
                [txt("طول ورق", "Plate length"), "Lp", _fmt(p.plate_length_mm), "mm"],
                [txt("ضخامت ورق", "Plate thickness"), "tp", _fmt(p.plate_thickness_mm), "mm"],
                [txt("Fy ورق", "Plate yield strength"), "Fyp", _fmt(result.plate_fy_mpa), "MPa"],
                [txt("Fu ورق", "Plate tensile strength"), "Fup", _fmt(result.plate_fu_mpa), "MPa"],
            ]
        )
    if b:
        _make_table(doc,
            headers=[txt("پارامتر", "Property"), txt("نماد", "Symbol"), txt("مقدار", "Value"), txt("واحد", "Unit")],
            rows=[
                [txt("نوع پیچ", "Bolt type"), "—", b.bolt_type.value, "—"],
                [txt("قطر پیچ", "Bolt diameter"), "db", _fmt(b.bolt_diameter_mm), "mm"],
                [txt("تعداد پیچ", "Number of bolts"), "n", str(b.num_bolts), "—"],
                [txt("نیروی هر پیچ", "Bolt force"), "Vu", _fmt(b.bolt_force_n / 1000.0), "kN"],
                [txt("ظرفیت هر پیچ", "Bolt capacity"), "φRn", _fmt(b.bolt_capacity_n / 1000.0), "kN"],
            ]
        )
    if w:
        _make_table(doc,
            headers=[txt("پارامتر", "Property"), txt("نماد", "Symbol"), txt("مقدار", "Value"), txt("واحد", "Unit")],
            rows=[
                [txt("سایز جوش", "Weld size"), "w", _fmt(w.weld_size_mm), "mm"],
                [txt("طول موثر جوش", "Effective weld length"), "Lw", _fmt(w.weld_length_mm), "mm"],
                [txt("ظرفیت جوش", "Weld capacity"), "φRn", _fmt(w.capacity_n / 1000.0), "kN"],
            ]
        )

    _heading(doc, txt("3. روند طراحی", "3. Design Procedure"), level=2)

    _heading(doc, txt("گام 1 — لنگر پلاستیک مورد انتظار", "Step 1 — Expected Plastic Moment"), level=3)
    _body(doc, txt("محاسبه ضریب سخت‌شوندگی و لنگر مورد انتظار تیر:", "Calculate strain hardening factor and expected beam moment:"), rtl=is_fa)
    _formula_line(doc, "C_pr", f"min((Fy + Fu)/(2×Fy), 1.2) = min(({_fmt(result.fy_beam_mpa)} + {_fmt(result.fu_beam_mpa)})/(2×{_fmt(result.fy_beam_mpa)}), 1.2)", _fmt(cpr), rtl=is_fa)
    _formula_line(doc, "M_pr", f"C_pr × R_y × F_y × Z_x = {_fmt(cpr)} × {_fmt(ry)} × {_fmt(result.fy_beam_mpa)} × {_fmt(zx_cm3)}", _fmt(mpr_knm), "kN·m", rtl=is_fa)

    _heading(doc, txt("گام 2 — نیروی کششی ورق سپری", "Step 2 — Flange Plate Force"), level=3)
    _formula_line(doc, "F_f", f"M_pr / d = {_fmt(mpr_knm * 1000.0)} / {_fmt(result.beam_depth_mm)}", _fmt(flange_force_kn), "kN", rtl=is_fa)

    if b:
        _heading(doc, txt("گام 3 — طراحی پیچ‌ها", "Step 3 — Bolt Design"), level=3)
        _formula_line(doc, "V_u", f"F_f / n = {_fmt(flange_force_kn)} / {b.num_bolts}", _fmt(b.bolt_force_n / 1000.0), "kN/bolt", rtl=is_fa)
        _formula_line(doc, "φR_n", txt("ظرفیت کششی هر پیچ", "Design strength per bolt"), _fmt(b.bolt_capacity_n / 1000.0), "kN/bolt", rtl=is_fa)
        _check_line(doc, txt("کنترل پیچ", "Bolt strength check"), _fmt(b.bolt_force_n / 1000.0), _fmt(b.bolt_capacity_n / 1000.0), b.is_adequate, b.utilization_ratio, is_fa)

    if p:
        ay = p.plate_width_mm * p.plate_thickness_mm
        py_cap = result.flange_force_n / p.yield_check_ratio / 1000.0 if p.yield_check_ratio > 0 else 0.0
        pr_cap = result.flange_force_n / p.rupture_check_ratio / 1000.0 if p.rupture_check_ratio > 0 else 0.0
        pb_cap = result.flange_force_n / p.block_shear_ratio / 1000.0 if p.block_shear_ratio > 0 else 0.0

        _heading(doc, txt("گام 4 — طراحی و کنترل ورق سپری", "Step 4 — Flange Plate Design and Checks"), level=3)
        _formula_line(doc, "A_g", f"b × t = {_fmt(p.plate_width_mm)} × {_fmt(p.plate_thickness_mm)}", _fmt(ay), "mm²", rtl=is_fa)
        _formula_line(doc, "t_req", txt("ضخامت موردنیاز محاسبه‌شده", "Calculated required thickness"), _fmt(p.required_thickness_mm), "mm", rtl=is_fa)
        _check_line(doc, txt("ضخامت ورق", "Plate thickness"), _fmt(p.required_thickness_mm), _fmt(p.plate_thickness_mm), p.plate_thickness_mm >= p.required_thickness_mm, None, is_fa)
        _formula_line(doc, "φP_n,y", txt("ظرفیت تسلیم ورق", "Plate yielding capacity"), _fmt(py_cap), "kN", rtl=is_fa)
        _check_line(doc, txt("کنترل تسلیم", "Yielding check"), _fmt(flange_force_kn), _fmt(py_cap), p.yield_check_ratio <= 1.0, p.yield_check_ratio, is_fa)
        _formula_line(doc, "φP_n,r", txt("ظرفیت گسیختگی ورق", "Plate rupture capacity"), _fmt(pr_cap), "kN", rtl=is_fa)
        _check_line(doc, txt("کنترل گسیختگی", "Rupture check"), _fmt(flange_force_kn), _fmt(pr_cap), p.rupture_check_ratio <= 1.0, p.rupture_check_ratio, is_fa)
        _formula_line(doc, "φP_n,bs", txt("ظرفیت بلوک برشی", "Block shear capacity"), _fmt(pb_cap), "kN", rtl=is_fa)
        _check_line(doc, txt("کنترل بلوک برشی", "Block shear check"), _fmt(flange_force_kn), _fmt(pb_cap), p.block_shear_ratio <= 1.0, p.block_shear_ratio, is_fa)

    if w:
        _heading(doc, txt("گام 5 — طراحی جوش", "Step 5 — Weld Design"), level=3)
        _formula_line(doc, "w_req", txt("سایز موردنیاز جوش", "Required weld size"), _fmt(w.required_size_mm), "mm", rtl=is_fa)
        _formula_line(doc, "φR_n", txt("ظرفیت جوش", "Weld design strength"), _fmt(w.capacity_n / 1000.0), "kN", rtl=is_fa)
        _check_line(doc, txt("کنترل جوش", "Weld strength check"), _fmt(flange_force_kn), _fmt(w.capacity_n / 1000.0), w.is_adequate, w.utilization_ratio, is_fa)

    _heading(doc, txt("گام 6 — کنترل لرزه‌ای و اضافه مقاومت", "Step 6 — Seismic Overstrength Check"), level=3)
    _formula_line(doc, "Ω_conn", txt("نسبت ظرفیت اتصال به 1.2M_pr", "Connection capacity ratio to 1.2M_pr"), _fmt(result.overstrength_ratio), "—", rtl=is_fa)
    _check_line(doc, txt("کنترل اضافه مقاومت", "Overstrength check"), "1.00", _fmt(result.overstrength_ratio), result.overstrength_ratio >= 1.0, result.overstrength_ratio, is_fa)

    _heading(doc, txt("4. خلاصه کنترل‌ها", "4. Design Checks Summary"), level=2)
    check_rows = [
        [txt("لنگر مورد انتظار", "Expected plastic moment"), _fmt(mpr_knm), txt("محاسبه شد", "Calculated"), "✓  OK"],
        [txt("کنترل پیچ", "Bolt strength"), _fmt(b.bolt_force_n / 1000.0) if b else "—", _fmt(b.bolt_capacity_n / 1000.0) if b else "—", "✓  OK" if (b and b.is_adequate) else "✗  FAIL"],
        [txt("کنترل تسلیم ورق", "Plate yielding"), _fmt(flange_force_kn) if p else "—", _fmt(py_cap) if p else "—", "✓  OK" if (p and p.yield_check_ratio <= 1.0) else "✗  FAIL"],
        [txt("کنترل گسیختگی ورق", "Plate rupture"), _fmt(flange_force_kn) if p else "—", _fmt(pr_cap) if p else "—", "✓  OK" if (p and p.rupture_check_ratio <= 1.0) else "✗  FAIL"],
        [txt("بلوک برشی", "Block shear"), _fmt(flange_force_kn) if p else "—", _fmt(pb_cap) if p else "—", "✓  OK" if (p and p.block_shear_ratio <= 1.0) else "✗  FAIL"],
        [txt("کنترل جوش", "Weld strength"), _fmt(flange_force_kn) if w else "—", _fmt(w.capacity_n / 1000.0) if w else "—", "✓  OK" if (w and w.is_adequate) else "✗  FAIL"],
        [txt("اضافه مقاومت", "Overstrength"), "1.00", _fmt(result.overstrength_ratio), "✓  OK" if result.overstrength_ratio >= 1.0 else "✗  FAIL"],
    ]
    summary_tbl = doc.add_table(rows=1 + len(check_rows), cols=4)
    summary_tbl.style = "Table Grid"
    summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(summary_tbl, [txt("کنترل", "Check"), txt("نیاز / تقاضا", "Demand"), txt("ظرفیت / وضعیت", "Capacity / Status"), txt("نتیجه", "Result")])
    for r_i, row_data in enumerate(check_rows):
        row = summary_tbl.rows[r_i + 1]
        for c_i, val in enumerate(row_data):
            cell = row.cells[c_i]
            pgr = cell.paragraphs[0]
            pgr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pgr.add_run(str(val))
            run.font.size = Pt(9)
            if c_i == 3:
                ok = "OK" in str(val)
                run.bold = True
                run.font.color.rgb = _OK if ok else _FAIL
    doc.add_paragraph()

    _heading(doc, txt("5. جمع‌بندی", "5. Summary"), level=2)
    _body(doc, result.summary, is_fa)

    final_p = doc.add_paragraph()
    final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_run = final_p.add_run(
        txt(
            f"✓ اتصال برای {std_name} قابل قبول است." if result.is_valid else f"✗ اتصال برای {std_name} نیاز به بازنگری دارد.",
            f"✓ Connection is adequate for {std_name}." if result.is_valid else f"✗ Connection needs revision for {std_name}.",
        )
    )
    final_run.bold = True
    final_run.font.size = Pt(12)
    final_run.font.color.rgb = _OK if result.is_valid else _FAIL

    if not output_path:
        output_path = "BFP_Connection_Report.docx"
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
